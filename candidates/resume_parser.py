# candidates/resume_parser.py

import logging
import mimetypes
import re
from pathlib import Path
from typing import Optional, Dict, Any, List

import numpy as np
import pypdfium2 as pdfium
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)

# ============================================================
#  OPTIONAL PADDLE OCR LOADING (SAFE)
# ============================================================

def load_paddle_ocr():
    """
    Safely load PaddleOCR ONLY if installed.
    If import fails → return None and the system will fallback to Tesseract or nothing.
    """
    try:
        from paddleocr import PaddleOCR
        return PaddleOCR(use_angle_cls=True, lang="en")
    except Exception as e:
        logger.warning(f"[Parser] PaddleOCR not available: {e}")
        return None


# Lazy-loaded global
PADDLE = None


# ============================================================
#  MIME / TYPE DETECTION
# ============================================================

def detect_mime_type(path: str) -> Optional[str]:
    path_obj = Path(path)
    mime, _ = mimetypes.guess_type(str(path_obj))
    if mime:
        return mime.lower()

    ext = path_obj.suffix.lower()
    if ext == ".pdf":
        return "application/pdf"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ext in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp"}:
        return f"image/{ext.lstrip('.')}"
    return None


# ============================================================
#  PDF PARSING (TEXT + SAFE OCR)
# ============================================================

def parse_pdf_text(path: str) -> str:
    doc = pdfium.PdfDocument(path)
    pages = []
    for page in doc:
        tp = page.get_textpage()
        txt = tp.get_text_range()
        if txt:
            pages.append(txt)
    return "\n".join(pages).strip()


def parse_pdf_ocr(path: str) -> str:
    global PADDLE
    if PADDLE is None:
        PADDLE = load_paddle_ocr()

    if PADDLE is None:
        logger.warning("[Parser] PaddleOCR unavailable → OCR skipped.")
        return ""

    doc = pdfium.PdfDocument(path)
    ocr_lines = []

    for page in doc:
        bitmap = page.render(scale=2.0)
        pil_img = Image.frombytes("RGB", (bitmap.width, bitmap.height), bitmap.to_bytes())
        np_img = np.array(pil_img)

        result = PADDLE.ocr(np_img, cls=True)
        if not result:
            continue

        for block in result:
            for line in block:
                text_line = line[1][0]
                if text_line:
                    ocr_lines.append(text_line)

    return "\n".join(ocr_lines).strip()


# ============================================================
#  DOCX PARSING
# ============================================================

def parse_docx_text(path: str) -> str:
    doc = Document(path)
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n".join(lines).strip()


# ============================================================
#  IMAGE OCR
# ============================================================

def parse_image_text(path: str) -> str:
    global PADDLE
    if PADDLE is None:
        PADDLE = load_paddle_ocr()

    if PADDLE is None:
        logger.warning("[Parser] PaddleOCR unavailable → image OCR skipped.")
        return ""

    result = PADDLE.ocr(path, cls=True)
    lines = []

    if not result:
        return ""

    for block in result:
        for line in block:
            text_line = line[1][0]
            if text_line:
                lines.append(text_line)

    return "\n".join(lines).strip()


# ============================================================
#  FIELD EXTRACTORS (unchanged)
# ============================================================

def extract_email(text: str) -> Optional[str]:
    m = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return m.group(0) if m else None


def extract_linkedin_username(text: str) -> Optional[str]:
    pattern = r"(linkedin)\s*[:\-•]\s*([A-Za-z0-9 ._'\-]+)"
    matches = re.findall(pattern, text, flags=re.IGNORECASE)
    if not matches:
        return None
    username = matches[0][1].strip()
    username = re.split(
        r"(portfolio|github|mail|email|phone)", username, flags=re.IGNORECASE
    )[0].strip()
    return username or None


def extract_name(text: str) -> Optional[str]:
    lines = text.strip().split("\n")[:10]
    banned = ["resume", "curriculum", "vitae", "summary", "education", "profile"]
    for line in lines:
        clean = line.strip()
        if not clean:
            continue
        lower = clean.lower()
        if "@" in clean or "linkedin.com" in lower:
            continue
        if any(b in lower for b in banned):
            continue
        if re.match(r"^[A-Za-z ,.\-']+$", clean) and 1 <= len(clean.split()) <= 5:
            return clean
    return None


SKILL_LIST = [
    "python", "java", "c++", "c", "sql", "nosql", "mysql", "postgresql", "mongodb",
    "aws", "azure", "gcp", "docker", "kubernetes", "pytorch", "tensorflow",
    "scikit-learn", "react", "django", "flask", "fastapi", "git", "linux",
    "javascript", "html", "css", "nlp", "machine learning", "deep learning",
    "computer vision", "tableau", "power bi"
]


def extract_skills(text: str) -> List[str]:
    tl = text.lower()
    found = {s for s in SKILL_LIST if re.search(rf"\b{s}\b", tl)}
    return sorted(found)


def extract_work_experience(text: str) -> Optional[str]:
    pattern = (
        r"(work experience|experience|professional experience|employment)"
        r"([\s\S]+?)(education|projects|skills|certifications|summary|$)"
    )
    match = re.search(pattern, text, flags=re.IGNORECASE)
    return match.group(2).strip() if match else None


def extract_rest(text: str, work_exp: Optional[str]) -> str:
    if not work_exp:
        return text.strip()
    idx = text.lower().find(work_exp.lower())
    if idx == -1:
        return text.strip()
    return text[idx + len(work_exp):].strip()


def parse_information(text: str) -> Dict[str, Any]:
    info = {
        "name": extract_name(text),
        "email": extract_email(text),
        "linkedin": extract_linkedin_username(text),
        "skills": extract_skills(text),
    }
    work_exp = extract_work_experience(text)
    info["work_experience"] = work_exp
    info["rest"] = extract_rest(text, work_exp)
    return info


# ============================================================
#  ENTRY POINT
# ============================================================

def _extract_text_by_type(path: str) -> str:
    mime = detect_mime_type(path)

    if mime == "application/pdf":
        text = parse_pdf_text(path)
        if len(text.strip()) < 50:
            text = parse_pdf_ocr(path)
        return text

    if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return parse_docx_text(path)

    if mime and mime.startswith("image/"):
        return parse_image_text(path)

    raise ValueError(f"Unsupported file format: {path}")


def parse_resume(file_path: str) -> Dict[str, Any]:
    logger.info(f"[Parser] Parsing: {file_path}")
    text = _extract_text_by_type(file_path)
    return parse_information(text)
